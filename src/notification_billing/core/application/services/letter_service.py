import io

from docx import Document
from docx.text.paragraph import Paragraph


class CordialLetterService:
    """
    Serviço para gerar uma carta em formato .docx a partir de um template,
    preenchendo placeholders com dados de um contexto.
    """

    def __init__(self, default_template_path: str):
        self.default_template_path = default_template_path

    def _replace_text_in_paragraph(self, paragraph: Paragraph, context: dict):
        if "{" not in paragraph.text:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        for key, value in context.items():
            full_text = full_text.replace(f"{{{key}}}", str(value))

        if paragraph.runs:
            original_run = paragraph.runs[0]
            original_run.text = full_text
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""

    def generate_letter(self, context: dict, template_path: str | None = None) -> io.BytesIO:
        """
        Carrega o template (padrão ou informado) e substitui os placeholders
        em todas as partes do documento.
        """
        docx_path = template_path or self.default_template_path
        document = Document(docx_path)

        for p in document.paragraphs:
            self._replace_text_in_paragraph(p, context)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_text_in_paragraph(p, context)

        for section in document.sections:
            if section.header:
                for p in section.header.paragraphs:
                    self._replace_text_in_paragraph(p, context)
            if section.footer:
                for p in section.footer.paragraphs:
                    self._replace_text_in_paragraph(p, context)

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream
